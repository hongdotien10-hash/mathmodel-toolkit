"""DOCX 题目解析器。"""

from pathlib import Path
from typing import Optional


class DocxReader:
    """Word 文档读取器。

    Usage::

        reader = DocxReader()
        text = reader.read("赛题.docx")
        tables = reader.extract_tables("赛题.docx")
    """

    def read(self, path: str | Path) -> str:
        """读取 DOCX 文本内容（含段落和表格文本）。

        Args:
            path: DOCX 文件路径

        Returns:
            str: 提取的文本
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("需要安装 python-docx: pip install python-docx")

        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {path}")

        doc = Document(str(path))
        parts = []

        # 段落文本
        for para in doc.paragraphs:
            if para.text.strip():
                # 识别标题样式
                if para.style and para.style.name and "Heading" in para.style.name:
                    parts.append(f"\n## {para.text.strip()}")
                else:
                    parts.append(para.text.strip())

        # 表格文本
        for i, table in enumerate(doc.tables):
            parts.append(f"\n[表格 {i + 1}]")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(" | ".join(cells))

        return "\n\n".join(parts)

    def extract_tables(self, path: str | Path) -> list[list[list[str]]]:
        """提取 DOCX 中的所有表格。

        Args:
            path: DOCX 文件路径

        Returns:
            list: 表格列表，每个表格是二维文本数组
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("需要安装 python-docx: pip install python-docx")

        doc = Document(str(path))
        tables = []
        for table in doc.tables:
            tbl = []
            for row in table.rows:
                tbl.append([cell.text.strip() for cell in row.cells])
            tables.append(tbl)
        return tables

    def get_metadata(self, path: str | Path) -> dict:
        """获取 DOCX 文档信息。"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("需要安装 python-docx: pip install python-docx")

        doc = Document(str(path))
        return {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "sections": len(doc.sections),
        }
