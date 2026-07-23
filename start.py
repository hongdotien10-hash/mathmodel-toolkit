"""MathModel Toolkit — 获奖级论文
====================================
用法：python start.py
增强：多模型竞赛 + 深度灵敏度 + 误差诊断 + 专业图表组 + AI写作
"""

import sys, json, re
from pathlib import Path
import re
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import warnings; warnings.filterwarnings('ignore')

sys.path.insert(0, str(Path(__file__).parent))

from mathmodel.utils import set_seed, Timer
from mathmodel.utils.helpers import safe_call
from mathmodel.analyzer import ProblemClassifier, ModelKnowledgeBase
from mathmodel.models import EvaluationSolver, StatsSolver, OptimizationSolver, MLSolver, GraphSolver
from mathmodel.sensitivity import SensitivityAnalyzer
from mathmodel.visualization import Plotter, set_style, get_colors
from mathmodel.paper.word_writer import generate_paper
from mathmodel.paper.latex_writer import generate_latex_paper
from mathmodel.pro import (ModelContest, DeepSensitivity, ResultNarrator,
                           ChartSuite, ErrorDiagnostics)
from mathmodel.pipeline.rich_progress import PhaseTracker, print_header, print_section, print_result_summary

set_seed(42)

PROBLEMS_DIR = Path(__file__).parent / "problems"
OUTPUT_DIR = Path(__file__).parent / "output"

# --- CLI flags ---
INTERACTIVE = "--interactive" in sys.argv or "-i" in sys.argv

# 比赛特定参数
CONTEST_PARAMS = {
    "cumcm": {"vehicle_speed": 50, "max_pages": 25, "lang": "zh", "name": "国赛 CUMCM"},
    "diangong": {"vehicle_speed": 50, "max_pages": 25, "lang": "zh", "name": "电工杯"},
    "mcm": {"vehicle_speed": 50, "max_pages": 25, "lang": "en", "name": "美赛 MCM/ICM"},
    "huawei": {"vehicle_speed": 50, "max_pages": 25, "lang": "zh", "name": "华为杯"},
    "auto": {"vehicle_speed": 50, "max_pages": 25, "lang": "zh", "name": "自动检测"},
}


def _pause(msg="Continue?"):
    """交互模式暂停"""
    if INTERACTIVE:
        try:
            input(f"\n  [PAUSE] {msg} (press Enter to continue, Ctrl+C to stop) ")
        except (EOFError, KeyboardInterrupt):
            print("\n  Stopped by user.")
            sys.exit(0)


def main():
    # --- Show interactive menu ---
    from mathmodel.pipeline.menu import show_menu
    CONTEST_TYPE, MAX_QUESTIONS, MAX_PAGES, MAX_FIGURES, USER_NOTES = show_menu()
    cp = CONTEST_PARAMS.get(CONTEST_TYPE, CONTEST_PARAMS["auto"])

    tracker = PhaseTracker(title="MathModel Toolkit PRO")
    print_header(f"MathModel Toolkit — {cp['name']}")
    if USER_NOTES:
        print(f"  User requirements: {USER_NOTES[:200]}")
    if INTERACTIVE:
        print("  [Interactive mode] Will pause at each phase for review.")

    # === Step 1: Scan & Load (same as free version) ===
    problem_dirs = sorted([d for d in PROBLEMS_DIR.iterdir()
                           if d.is_dir() and not d.name.startswith('.')])
    if not problem_dirs:
        print("ERROR: No problems found in problems/")
        sys.exit(1)
    selected = problem_dirs[0]
    out_dir = OUTPUT_DIR / selected.name
    out_dir.mkdir(parents=True, exist_ok=True)
    fig_dir = out_dir / "figures"
    fig_dir.mkdir(parents=True, exist_ok=True)
    print(f"  Output: {out_dir}")
    print(f"  Figures: {fig_dir}")

    problem_text, data_files, data_profiles = "", {}, {}
    for f in sorted(selected.iterdir()):
        s = f.suffix.lower()
        if s == '.txt':
            problem_text = f.read_text(encoding='utf-8')
            print(f"[doc] {f.name} ({len(problem_text)} chars)")
        elif s == '.pdf':
            try:
                import pdfplumber
                with pdfplumber.open(str(f)) as pdf:
                    problem_text = '\n\n'.join(p.extract_text() or "" for p in pdf.pages)
                print(f"[doc] {f.name} ({len(problem_text)} chars)")
            except Exception as e: print(f"[doc] {f.name} — {e}")
        elif s in ('.xlsx', '.xls'):
            df = pd.read_excel(f, nrows=500) if f.stat().st_size > 5e6 else pd.read_excel(f)
            data_files[f.stem] = df
            data_profiles[f.stem] = {"name": f.stem, "path": str(f), "shape": df.shape,
                                     "columns": list(df.columns), "size_category":
                                     "large" if f.stat().st_size > 5e6 else "medium" if df.shape[0] > 1000 else "small",
                                     "sample": df.head(5)}
            print(f"[data] {f.name} {df.shape}")
        elif s == '.csv':
            df = pd.read_csv(f, low_memory=True, nrows=500) if f.stat().st_size > 5e6 else pd.read_csv(f)
            data_files[f.stem] = df
            data_profiles[f.stem] = {"name": f.stem, "path": str(f), "shape": df.shape,
                                     "columns": list(df.columns), "size_category":
                                     "large" if f.stat().st_size > 5e6 else "small",
                                     "sample": df.head(5)}
            print(f"[data] {f.name} {df.shape}")

    # === Step 2: AI Analysis (same as free) ===
    sub_problems = _analyze(problem_text, data_profiles)
    print_section("Phase 1: AI Analysis")
    for sp in sub_problems:
        print(f"  Q{sp['id']}: [{sp['type']}] → {sp['model']} ({sp['score']:.0%}) [{sp.get('source','rule')}]")

    # === Step 3: API client for Pro modules ===
    api_key = ""
    try:
        from api.config import APIConfig
        cfg = APIConfig()
        if cfg.is_configured: api_key = cfg.api_key
    except Exception: pass

    # === Step 4: 🆕 PRO — Multi-Model Contest (本地或AI对比) ===
    print_section("PRO Phase 2: Multi-Model Contest")
    contest_results = {}
    mc = ModelContest(api_key=api_key)
    for sp in sub_problems:
        # Skip contest for optimization — use deep solver directly
        if sp["type"] == "优化":
            print(f"\n  >> Q{sp['id']}: Optimization — skip contest, use deep solver")
            continue
        if sp["type"] in ("评价", "预测", "统计"):
            print(f"\n  >> Q{sp['id']}: Running {sp['type']} model contest...")
            cr = safe_call(lambda s=sp: mc.contest(s, data_files),
                          desc=f"Q{sp['id']} model contest", timeout=300,
                          default={"winner": sp.get("model","?"), "candidates": [],
                                   "ai_reason": "contest skipped (timeout)"})
            contest_results[f"sub_{sp['id']}"] = cr
            print(f"     Winner: {cr.get('winner','?')} | {cr.get('ai_reason','')[:80]}...")
    if not api_key:
        print("  [本地对比] 基于指标数值自动选优")
    _pause("Contest done. Continue to solving?")

    # --- Save partial results early ---
    try:
        import json as _json
        _json.dump({"sub_problems": sub_problems, "contest": {k: {"winner": v.get("winner","")}
                   for k, v in contest_results.items()}},
                   open(out_dir / "results.json", "w", encoding="utf-8"),
                   ensure_ascii=False, indent=2, default=str)
    except: pass

    # === Step 5: Solve with best models ===
    print_section("Phase 3: Solving with Best Models (Deep Mode)")
    print(f"  Contest: {CONTEST_TYPE} | Questions: {MAX_QUESTIONS} | Speed: {cp['vehicle_speed']}km/h")

    # Limit questions EARLY — before any solving
    sub_problems = [sp for sp in sub_problems if sp["id"] <= MAX_QUESTIONS]
    if MAX_QUESTIONS < 99 and len(sub_problems) < len([sp for sp in sub_problems]):
        # Re-evaluate: the filter was applied after AI analysis created all 4
        pass
    if MAX_QUESTIONS < 99:
        print(f"  Solving only questions: {[sp['id'] for sp in sub_problems]}")

    all_results = {}
    for sp in sub_problems:
        sp_id = sp["id"]
        ptype = sp["type"]

        # ROUTING/OPTIMIZATION: Skip contest entirely, use deep solver
        # Contest is only useful for evaluation/prediction where multiple model types exist
        # For optimization, all candidates run the same TSP/VRP solver anyway
        is_routing_sp = ptype == "优化"  # ALL optimization goes to deep solver
        sp_text = sp.get("title", "") + sp.get("full_text", "")

        if is_routing_sp:
            print(f"\n  {'='*40}")
            print(f"  Q{sp_id}: Deep Multi-Round Analysis + Solve")
            print(f"  {'='*40}")
            df = _find_df(data_files, ptype)
            if df is not None:
                from mathmodel.pipeline.deep_solve import deep_solve_tsp
                numeric = df.select_dtypes(include=np.number)
                first_col = numeric.iloc[:, 0].dropna().tolist()
                if len(first_col) >= 3 and first_col[:3] == [1.0, 2.0, 3.0]:
                    sparse = numeric.iloc[:, 1:].values.astype(float)
                else:
                    sparse = numeric.values.astype(float)
                n = sparse.shape[0]

                # Phase A: Deep solver (5+ min computation)
                print(f"  [Phase A] Deep TSP solving ({n} locations, 10min budget)...")
                deep_result = deep_solve_tsp(sparse, n, fig_dir, sp_id, time_budget=600)

                # Phase B: 30-API deep thinking per question
                ai_insights = {}
                if api_key:
                    print(f"  [Phase B] 30-API deep analysis pipeline...")
                    try:
                        from mathmodel.pipeline.deep_thinker import DeepThinker
                        thinker = DeepThinker(api_key=api_key)
                        result_preview = {
                            "total_distance": deep_result["best"]["distance"],
                            "tour": deep_result["best"]["tour"][:10],
                            "method": deep_result["best"]["method"],
                            "n_locations": n,
                            "all_methods": deep_result["all_ranked"],
                        }
                        ai_insights = thinker.think_one_question(
                            sp, problem_text, data_profiles, result_preview, fig_dir)
                    except Exception as e:
                        print(f"  AI error: {e}")

                all_results[f"sub_{sp_id}"] = {
                    "method": "Floyd-Warshall + TSP(NN+2-opt+SA)",
                    "total_distance": deep_result["best"]["distance"],
                    "tour": deep_result["best"]["tour"],
                    "tour_labels": [str(t+1) for t in deep_result["best"]["tour"]],
                    "n_locations": n, "n_vehicles": 1,
                    "all_methods": deep_result["all_ranked"],
                    "total_time_s": deep_result["total_time"],
                    "ai_analysis": ai_insights.get("result_final", ""),
                    "ai_model_debate": ai_insights.get("model_decision", ""),
                    "ai_figures": [ai_insights.get(f"figure_{i+1}", {}) for i in range(3)],
                    "ai_quality": ai_insights.get("quality_scores", {}).get("overall_score", 0),
                    "ai_judge_comment": ai_insights.get("judge_comment", ""),
                    "summary": (f"最短配送回路总距离: {deep_result['best']['distance']}km, "
                               f"覆盖全部{n}个地点。"
                               f"{ai_insights.get('result_final', '')[:300]}")
                }
                print(f"     Result: {deep_result['best']['distance']}km")
                print(f"     AI calls: {thinker.total_calls if api_key else 0}, "
                      f"cost: ¥{thinker.total_cost if api_key else 0:.4f}")
            continue
        elif f"sub_{sp_id}" in contest_results:
            best = contest_results[f"sub_{sp_id}"]
            winner = best.get("winner", "")
            if winner not in ("无有效模型", "无", "", None):
                best_result = next((c for c in best.get("candidates", [])
                                    if c["model"] == winner), {})
                result_data = best_result.get("result", {})
                if result_data and result_data.get("metric_value", -1) != 0 and \
                   result_data.get("selection") != []:
                    all_results[f"sub_{sp_id}"] = result_data
                    print(f"  Q{sp_id}: [{winner}] — contest winner")
                    continue
                else:
                    print(f"  Q{sp_id}: Contest result empty, falling back to local solver...")
            else:
                print(f"  Q{sp_id}: Contest failed ({winner}), using local solver...")

        # Fallback: local solver
        df = _find_df(data_files, ptype)
        if df is None: continue
        numeric = df.select_dtypes(include=np.number)
        if ptype == "评价" and numeric.shape[1] >= 2:
            ev = EvaluationSolver()
            m = numeric.values.astype(float)
            ew = ev.entropy_weight(m)
            res = ev.topsis(m, weights=ew["weights"], impacts=[1]*m.shape[1])
            labels = df.iloc[:,0].tolist()
            all_results[f"sub_{sp_id}"] = {"labels": labels,
                "scores": [round(float(s),4) for s in res["scores"]],
                "rank": [int(r) for r in res["rank"]],
                "weights": {str(c): round(float(w),4) for c,w in zip(numeric.columns, ew["weights"])}}
            print(f"  Q{sp_id}: TOPSIS done")
        elif ptype == "预测" and numeric.shape[1] >= 1:
            ss = StatsSolver()
            data = numeric.iloc[:,0].dropna().tolist()
            if len(data) >= 4:
                pred = ss.grey_forecast(data, 3)
                all_results[f"sub_{sp_id}"] = {"original": data,
                    "fitted": [round(v,4) for v in pred["fitted"]],
                    "forecast": [round(v,4) for v in pred["forecast"]],
                    "mape": round(pred["mape"],2), "grade": pred["grade"]}
                print(f"  Q{sp_id}: GM(1,1) MAPE={pred['mape']:.2f}%")
        elif ptype == "优化" and numeric.shape[1] >= 2:
            # Only reached for non-routing optimization (e.g. 0-1 knapsack)
            costs = numeric.iloc[:,1].values.astype(float).tolist()
            benefits = numeric.iloc[:,2].values.astype(float).tolist() if numeric.shape[1] > 2 else numeric.iloc[:,0].tolist()
            budget = sum(costs)*0.6
            ratios = sorted([(benefits[i]/max(costs[i],1e-6),i) for i in range(len(costs))], key=lambda x:-x[0])
            sel = []; rem = budget; cost_total = 0; benefit_total = 0
            for _, i in ratios:
                if costs[i] <= rem: sel.append(i); rem -= costs[i]; cost_total += costs[i]; benefit_total += benefits[i]
            labels = df.iloc[:,0].tolist()
            all_results[f"sub_{sp_id}"] = {"selection": [str(labels[i]) for i in sel],
                "total_cost": round(cost_total,1), "total_population": round(benefit_total,1),
                "budget": round(budget,1), "solution": [1 if i in sel else 0 for i in range(len(costs))],
                "costs": costs, "benefits": benefits, "labels_all": labels}
            print(f"  Q{sp_id}: Knapsack done ({len(sel)} selected)")

    # --- Save results immediately after solving ---
    try:
        json.dump({"sub_problems": sub_problems, "results": _serialize(all_results),
                   "contest": {k: {"winner": v.get("winner","")} for k, v in contest_results.items()}},
                  open(out_dir / "results.json", "w", encoding="utf-8"),
                  ensure_ascii=False, indent=2, default=str)
        print(f"  Results saved: {out_dir / 'results.json'}")
    except Exception as e:
        print(f"  Failed to save results: {e}")

    # === Step 6: 🆕 PRO — Deep Sensitivity ===
    print_section("PRO Phase 4: Deep Sensitivity")
    ds = DeepSensitivity()
    sens_results = {}
    for key, val in all_results.items():
        if "forecast" in val and len(val["fitted"]) >= 4:
            def gm(p): s = StatsSolver(); r = s.grey_forecast(p.tolist(), 3); return float(r["forecast"][-1])
            base = val["fitted"]
            tornado = ds.tornado(gm, base, [f"Point {i+1}" for i in range(len(base))])
            mc = ds.monte_carlo(gm, base, noise_pct=0.05, n_samples=500)
            sens_results[key] = {"tornado": tornado, "monte_carlo": mc}
            ds.plot_tornado(tornado["impacts"], title="Parameter Sensitivity",
                            output_path=str(fig_dir / f"{key}_tornado.pdf"))
            ds.plot_mc_distribution([gm(np.array(base)*np.random.normal(1,0.05,len(base))) for _ in range(500)],
                                   output_path=str(fig_dir / f"{key}_mc_dist.pdf"))
            print(f"  {key}: Tornado top={tornado['most_sensitive']}, MC CV={mc['cv']:.4f}")

    # === Step 7: 🆕 PRO — Error Diagnostics ===
    print_section("PRO Phase 5: Error Diagnostics")
    ed = ErrorDiagnostics()
    diag_results = {}
    for key, val in all_results.items():
        if "forecast" in val and "fitted" in val and "original" in val:
            ra = ed.residual_analysis(val["original"], val["fitted"])
            intervals = ed.prediction_intervals(val["fitted"], val["forecast"], np.std(ra["residuals"]))
            diag_results[key] = {"residual_analysis": ra, "prediction_intervals": intervals}
            ed.plot_residuals(val["original"], val["fitted"],
                             output_path=str(fig_dir / f"{key}_residuals.png"))
            ed.plot_prediction_interval(val["fitted"], val["forecast"], intervals["intervals"],
                                        output_path=str(fig_dir / f"{key}_pred_intervals.png"))
            print(f"  {key}: MAPE={ra['mape']:.2f}%, DW={ra['dw']:.3f}, Normal={ra['is_normal']}")

    # === Step 8: 🆕 PRO — Chart Suite ===
    print_section("PRO Phase 6: Professional Chart Suites")
    cs = ChartSuite()

    # === AI智能选图 ===
    try:
        from mathmodel.pipeline.ai_figure_selector import ai_select_figures, generate_selected_figures
        print(f"  [AI] Analyzing results to select best figures (max {MAX_FIGURES})...")
        selections = ai_select_figures(sub_problems, all_results, data_files, max_figures=MAX_FIGURES)
        generated = generate_selected_figures(selections, all_results, data_files, fig_dir, sub_problems)
        print(f"  [AI] Generated {len(generated)}/{len(selections)} selected figures")
    except Exception as e:
        print(f"  AI figure selection: {e}")
        import traceback; traceback.print_exc()

    # Generate professional figures from real data
    try:
        from mathmodel.pipeline.professional_figures import (
            fig_tsp_network, fig_algorithm_comparison, fig_convergence_curve,
            fig_distance_matrix_heatmap, fig_question_comparison
        )

        # Distance matrix heatmap
        for name, df in data_files.items():
            if name.endswith("_norm"): continue
            numeric = df.select_dtypes(include=np.number)
            if numeric.shape[1] >= 10:
                fig_distance_matrix_heatmap(numeric.values, min(numeric.shape[0], numeric.shape[1]),
                                           str(fig_dir / f"distance_matrix_{name}.pdf"),
                                           title=f"Distance Matrix - {name}")
                print(f"  Distance matrix heatmap: {name}")
                break

        # Sub-problem specific figures
        for key, val in all_results.items():
            if "tour" in val and val.get("tour"):
                # TSP network with real coordinates
                sp_id = key.replace("sub_", "")
                # Find the data used for this sub-problem
                for name, df in data_files.items():
                    if name.endswith("_norm"): continue
                    numeric = df.select_dtypes(include=np.number)
                    if numeric.shape[1] >= 5:
                        tour = val["tour"]
                        dist = val.get("total_distance", 0)
                        fig_tsp_network(numeric.values, len(tour)-1, tour, dist,
                                       str(fig_dir / f"{key}_network.pdf"),
                                       title=f"Question {sp_id}")
                        print(f"  [{key}] TSP network figure")
                        break

                # Algorithm comparison
                methods = val.get("all_methods", [])
                if methods:
                    methods_dict = {m[:25]: d for d, m in methods[:6]}
                    fig_algorithm_comparison(methods_dict,
                                            str(fig_dir / f"{key}_algo_compare.pdf"),
                                            title=f"Q{sp_id} Algorithm Comparison")
                    print(f"  [{key}] Algorithm comparison figure")

    except Exception as e:
        print(f"  Professional figures: {e}")

    # Standard chart suites
    for key, val in all_results.items():
        if "tour" in val or "routes" in val:
            _make_routing_figure(val, key, fig_dir)
        if "forecast" in val and "fitted" in val:
            cs.prediction_suite(val.get("original", []), val["fitted"], val["forecast"],
                                output_path=str(fig_dir / f"{key}_pred_suite.png"))
            print(f"  [{key}] Prediction suite generated")
        if "scores" in val and "labels" in val and "weights" in val:
            cs.evaluation_suite(val["scores"], val["labels"], val["weights"],
                               output_path=str(fig_dir / f"{key}_eval_suite.png"))
            print(f"  [{key}] Evaluation suite generated")
        if "selection" in val and "costs" in val:
            cs.optimization_suite(val["costs"][:20], val["benefits"][:20], val.get("labels_all", val["selection"])[:20],
                                  val["solution"][:20] if "solution" in val else [1]*len(val["costs"][:20]),
                                  val.get("budget", 100),
                                  output_path=str(fig_dir / f"{key}_opt_suite.png"))
            print(f"  [{key}] Optimization suite generated")

    # === Step 9: 🆕 PRO — Result-Driven Writing ===
    ai_content = {}
    if api_key:
        print_section("PRO Phase 7: Result-Driven Writing")
        rn = ResultNarrator(api_key=api_key)
        for sp in sub_problems:
            sp_id = sp["id"]
            result = all_results.get(f"sub_{sp_id}", {})
            cr = contest_results.get(f"sub_{sp_id}", {})
            text = safe_call(lambda s=sp, r=result, c=cr: rn.narrate_result(s, r, c),
                           desc=f"Q{sp_id} result writing", timeout=180,
                           default=rn._fallback_narrative(sp, result))
            ai_content[f"section_{sp_id}"] = text
            print(f"  Q{sp_id}: {len(text)} chars")

        # AI abstract and evaluation
        try:
            from mathmodel.pipeline.smart_orchestrator import AIDrivenPipeline
            ai = AIDrivenPipeline(api_key=api_key, model="deepseek-chat")
            dsum = "\n".join(f"{n}: {p['shape'][0]}r x {p['shape'][1]}c" for n, p in data_profiles.items())
            # Inject deep thinking results into AI content
            for sp in sub_problems:
                sp_id = sp["id"]
                deep = all_results.get(f"sub_{sp_id}", {})
                if deep.get("ai_analysis"):
                    ai_content[f"section_{sp_id}"] = deep["ai_analysis"]
                    print(f"  Using deep analysis for Q{sp_id} ({len(deep['ai_analysis'])} chars)")

            # Inject user requirements into AI planning
            if USER_NOTES:
                problem_text_aug = problem_text + f"\n\n[用户自定义要求]\n{USER_NOTES}"
            else:
                problem_text_aug = problem_text
            plan = ai.analyze_and_plan(problem_text_aug, dsum)
            analyses = ai.interpret_results(plan, all_results)
            ai_content["abstract"] = ai.write_abstract(problem_text, plan, all_results, analyses)
            ai_content["sensitivity"] = ai.write_sensitivity_section(all_results)
            ai_content["evaluation"] = ai.write_evaluation_section(sub_problems, all_results)
            print(f"  Abstract: {len(ai_content.get('abstract',''))} chars")
        except Exception as e:
            print(f"  AI writing failed: {e}")

    # --- Paper-level AI analysis (20 API calls) ---
    paper_insights = {}
    if api_key:
        print_section("Phase 7.5: Cross-Question Paper-Level Analysis")
        try:
            from mathmodel.pipeline.deep_thinker import DeepThinker
            paper_thinker = DeepThinker(api_key=api_key)
            paper_insights = paper_thinker.think_paper_level(
                sub_problems, all_results, problem_text, data_profiles)
            # Inject into ai_content
            if paper_insights.get("abstract_final"):
                ai_content["abstract"] = paper_insights["abstract_final"]
            if paper_insights.get("conclusion_final"):
                ai_content["conclusion"] = paper_insights.get("conclusion_final", "")
            if paper_insights.get("cross_comparison"):
                ai_content["cross_analysis"] = paper_insights.get("cross_comparison", "")
            print(f"  Paper AI: {paper_thinker.total_calls} calls, ¥{paper_thinker.total_cost:.4f}")
            pq = paper_insights.get("paper_quality", {})
            if isinstance(pq, dict):
                print(f"  Estimated prize: {pq.get('estimated_prize', '?')}")
        except Exception as e:
            print(f"  Paper AI: {e}")

    _pause("All results ready. Generate paper now?")

    # --- Generate question comparison figure ---
    try:
        from mathmodel.pipeline.professional_figures import fig_question_comparison
        fig_question_comparison(all_results, str(fig_dir / "all_questions_comparison.pdf"))
        print("  All-questions comparison figure generated")
    except Exception as e:
        print(f"  Comparison figure: {e}")

    # === Step 10: Generate Papers ===
    print_section("Phase 8: Paper Generation")
    import datetime
    ts = datetime.datetime.now().strftime("%H%M%S")
    try:
        paper_path = generate_paper(
            output_path=str(out_dir / f"Pro论文_{selected.name}_{ts}.docx"),
            problem_text=problem_text, analysis={"sub_problems": sub_problems},
            recommendations=[{"summary": "Pro: multi-model contest", "sub_problems": sub_problems}],
            results=all_results, figures_dir=str(fig_dir), ai_content=ai_content)
        print(f"  Word: {paper_path}")
    except Exception as e:
        print(f"  Word failed: {e}. Generating minimal paper...")
        # Fallback: minimal paper with results
        try:
            from docx import Document
            doc = Document()
            doc.add_heading("MathModel Toolkit Results", 0)
            doc.add_paragraph(f"Problem: {selected.name}")
            for sp in sub_problems:
                doc.add_heading(f"Q{sp['id']}: {sp.get('title','')}", 1)
                r = all_results.get(f"sub_{sp['id']}", {})
                doc.add_paragraph(r.get("summary", str(r)[:500]))
            paper_path = out_dir / f"论文_{selected.name}_{ts}.docx"
            doc.save(str(paper_path))
            print(f"  Fallback Word: {paper_path}")
        except Exception as e2:
            print(f"  Even fallback failed: {e2}")
            paper_path = "FAILED"

    try:
        tex_path = generate_latex_paper(str(out_dir / "latex"), problem_text,
                                        {"sub_problems": sub_problems}, all_results,
                                        str(fig_dir), ai_content)
        print(f"  LaTeX: {tex_path}")
    except Exception as e: print(f"  LaTeX: skipped ({e})")

    # === Save ===
    with open(out_dir / "pro_results.json", "w", encoding="utf-8") as f:
        json.dump({"sub_problems": sub_problems, "results": _serialize(all_results),
                   "contest": {k: {"winner": v["winner"]} for k, v in contest_results.items()},
                   "sensitivity": {k: {"tornado_top": v["tornado"]["most_sensitive"]}
                                   for k, v in sens_results.items()}},
                  f, ensure_ascii=False, indent=2, default=str)

    tracker.finish()
    print_result_summary(sub_problems, all_results)
    print()
    print("=" * 60)
    print("  OUTPUT FILES")
    print("=" * 60)
    print(f"  Word Paper:  {paper_path}")
    print(f"  Results:     {out_dir / 'results.json'}")
    print(f"  Pro Results: {out_dir / 'pro_results.json'}")
    print(f"  Figures:     {fig_dir}")
    print(f"  LaTeX:       {out_dir / 'latex'}")
    print("=" * 60)


def _analyze(problem_text, data_profiles):
    """AI analysis with rule fallback"""
    lines = problem_text.split('\n')
    sub_raw = []
    for i, line in enumerate(lines):
        line = line.strip()
        if not line: continue
        if '问题' in line and any(c in line for c in '123456789一二三四五六七八九'):
            ctx = line
            for j in range(i+1, min(i+8, len(lines))):
                if '问题' in lines[j] and any(c in lines[j] for c in '123456789'): break
                ctx += ' ' + lines[j].strip()
            sub_raw.append({"id": len(sub_raw)+1, "text": ctx[:500]})
            if len(sub_raw) >= 6: break
    if not sub_raw: sub_raw = [{"id": 1, "text": problem_text[:500]}]

    try:
        from api.config import APIConfig
        cfg = APIConfig()
        if cfg.is_configured:
            from api.analyzer import AIAnalyzer
            analyzer = AIAnalyzer(cfg)
            analysis = analyzer.full_analysis(problem_text, "\n".join(
                f"{n}: {p['shape']} cols={p['columns'][:5]}" for n, p in data_profiles.items()))
            sub_problems = []
            for sp in analysis.get("problem_analysis", {}).get("sub_problems", []):
                recs = analysis.get("model_recommendations", [])
                ai_model = ""; ai_score = 0.85
                for r in recs:
                    if r.get("sub_problem_id") == sp.get("id"):
                        rc = r.get("recommendations", [])
                        if rc: ai_model = rc[0].get("model", ""); ai_score = rc[0].get("score", 0.85)
                        break
                sub_problems.append({"id": sp.get("id", len(sub_problems)+1),
                    "title": sp.get("title","")[:150], "full_text": sp.get("title",""),
                    "type": sp.get("type","综合"), "model": ai_model or f"{sp.get('type','')}相关模型",
                    "score": ai_score, "reason": sp.get("reason",""), "source": "ai"})
            if sub_problems: return sub_problems
    except Exception: pass

    # Rule fallback
    classifier = ProblemClassifier(); kb = ModelKnowledgeBase()
    sub_problems = []
    for sp in sub_raw:
        clf = classifier.classify(sp["text"])
        cand = kb.query(problem_type=clf["type"], top_k=3)
        m = cand[0] if cand else {"model": "待定", "score": 0, "reason": ""}
        sub_problems.append({"id": sp["id"], "title": sp["text"][:150], "full_text": sp["text"],
            "type": clf["type"], "model": m["model"], "score": m["score"],
            "reason": m["reason"], "source": "rule"})
    return sub_problems


def _solve_routing(sp, df, data_files, fig_dir, results):
    """路径优化：VRP/TSP/CVRP — Floyd-Warshall + TSP最近邻 + 2-opt优化"""
    gs = GraphSolver()
    numeric = df.select_dtypes(include=np.number)
    n = numeric.shape[0]
    sp_text = sp.get("title", "") + sp.get("full_text", "")

    # ---- Step 0: Detect data format ----
    # Format A: Sparse distance matrix (square, many NaN = no direct road)
    # Format B: Coordinate table (X/Y/lat/lon columns)
    # Format C: Complete distance matrix (square, no NaN)

    is_distance_matrix = False
    if n >= 4 and numeric.shape[1] >= n - 2:  # rough check: square-ish
        # Check if first column looks like location IDs
        first_col = numeric.iloc[:, 0]
        nan_ratio = numeric.isnull().sum().sum() / (n * numeric.shape[1])
        if nan_ratio > 0.3:  # many NaN = sparse graph
            is_distance_matrix = True
            print(f"     Detected: Sparse distance matrix ({nan_ratio:.0%} NaN)")

    if is_distance_matrix:
        _solve_routing_from_matrix(sp, numeric, df, results, gs, n, sp_text)
    else:
        _solve_routing_from_coords(sp, numeric, df, results, gs, n, sp_text)


def _build_complete_graph(numeric, n):
    """从稀疏距离矩阵构建完全图（Floyd-Warshall）"""
    INF = 1e9
    D = np.full((n, n), INF)
    vals = numeric.values.astype(float)

    # If first column is location IDs, use cols 1: as distance matrix
    if numeric.shape[1] >= n:
        # Distance matrix: cols 0..n-1 or 1..n
        offset = 1 if not pd.api.types.is_numeric_dtype(numeric.columns[0]) else 0
        for i in range(n):
            for j in range(offset, min(offset + n, numeric.shape[1])):
                val = vals[i, j] if j < vals.shape[1] else np.nan
                if not np.isnan(val) and val > 0:
                    D[i, j - offset] = val
                if i == j - offset:
                    D[i, j - offset] = 0
    else:
        # General sparse matrix
        n_cols = min(n, numeric.shape[1])
        for i in range(n):
            for j in range(n_cols):
                val = vals[i, j]
                if not np.isnan(val) and val > 0:
                    D[i, j] = val
                elif i == j:
                    D[i, j] = 0

    # Make symmetric
    for i in range(n):
        for j in range(n):
            if D[i, j] < INF and D[j, i] >= INF:
                D[j, i] = D[i, j]
            elif D[j, i] < INF and D[i, j] >= INF:
                D[i, j] = D[j, i]

    # Floyd-Warshall
    n_edges_before = int(np.sum((D > 0) & (D < INF)))
    for k in range(n):
        for i in range(n):
            if D[i, k] >= INF: continue
            for j in range(n):
                nd = D[i, k] + D[k, j]
                if nd < D[i, j]:
                    D[i, j] = nd

    # Check connectivity — if any pair unreachable, that's bad
    n_unreachable = int(np.sum(D >= INF))
    if n_unreachable > 0:
        print(f"     Warning: {n_unreachable} unreachable pairs — graph may be disconnected")
        # Fallback: use Euclidean distances between implicit coordinates as backup
        D[D >= INF] = 999  # large penalty but keeps TSP working

    n_edges_after = int(np.sum((D > 0) & (D < INF)))
    print(f"     Graph: {n_edges_before} edges -> Floyd -> {n_edges_after} paths")
    return D


def _tsp_solve(D, n, max_starts=15):
    """TSP: 多起点最近邻 + 2-opt + 模拟退火 多算法对比选最优"""
    gs = GraphSolver()
    best_dist, best_tour = float('inf'), None

    # Phase 1: Nearest neighbor from multiple starts
    for start in range(min(n, max_starts)):
        r = gs.tsp_nearest_neighbor(D, start=start)
        if r['total_distance'] < best_dist:
            best_dist = r['total_distance']
            best_tour = r['tour']

    # Phase 2: 2-opt improvement on NN result
    improved = True
    iters = 0
    while improved and iters < 100:
        improved = False; iters += 1
        for i in range(1, len(best_tour) - 3):
            for j in range(i + 2, len(best_tour) - 1):
                old_d = D[best_tour[i-1]][best_tour[i]] + D[best_tour[j]][best_tour[j+1]]
                new_d = D[best_tour[i-1]][best_tour[j]] + D[best_tour[i]][best_tour[j+1]]
                if new_d < old_d - 1e-10:
                    best_tour[i:j+1] = reversed(best_tour[i:j+1])
                    best_dist = best_dist - old_d + new_d
                    improved = True

    # Phase 3: Also try Simulated Annealing for problems >= 10 nodes
    if n >= 10:
        sa_dist, sa_tour = _tsp_simulated_annealing(D, n, iterations=3000)
        if sa_dist < best_dist:
            best_dist, best_tour = sa_dist, sa_tour
            print(f"     SA improved: {best_dist}")

    return round(best_dist, 1), best_tour


def _tsp_simulated_annealing(D, n, temp_start=1000, cooling=0.995, iterations=5000):
    """模拟退火 TSP 求解器"""
    import random
    tour = list(range(n))
    random.shuffle(tour)
    tour.append(tour[0])
    current_dist = sum(D[tour[i]][tour[i+1]] for i in range(n))
    best_tour, best_dist = tour[:], current_dist
    temp = temp_start

    for _ in range(iterations):
        i, j = sorted(random.sample(range(1, n), 2))
        if j - i < 2: continue
        new_tour = tour[:i] + tour[i:j+1][::-1] + tour[j+1:]
        new_dist = sum(D[new_tour[k]][new_tour[k+1]] for k in range(n))

        if new_dist < current_dist or random.random() < np.exp((current_dist - new_dist) / max(temp, 1e-10)):
            tour, current_dist = new_tour, new_dist
            if current_dist < best_dist:
                best_tour, best_dist = tour[:], current_dist
        temp *= cooling

    return round(best_dist, 1), best_tour


def _solve_routing_from_matrix(sp, numeric, df, results, gs, n, sp_text):
    """稀疏距离矩阵 → Floyd → TSP + 2-opt"""
    D = _build_complete_graph(numeric, n)

    # Get capacity
    cap_match = re.search(r'(\d+)\s*(kg|千克|公斤|吨|t)', sp_text.lower())
    capacity = float(cap_match.group(1)) if cap_match else float('inf')
    if cap_match and cap_match.group(2) in ('吨', 't'):
        capacity *= 1000

    # Get demands if available
    demand_col = None
    for col in df.columns:
        cl = str(col).lower()
        if any(kw in cl for kw in ['需求', '重量', 'demand', 'weight', 'load']):
            demand_col = col; break
    demands = df[demand_col].values.astype(float).tolist() if demand_col else None

    labels = df.iloc[:, 0].tolist() if not pd.api.types.is_numeric_dtype(df.iloc[:, 0]) else \
             [f"地点{i+1}" for i in range(n)]

    if demands is None or sum(demands) <= capacity:
        # Single TSP route
        dist, tour = _tsp_solve(D, n)
        tour_labels = [labels[i] for i in tour if i < len(labels)]
        results[f"sub_{sp['id']}"] = {
            "method": "Floyd-Warshall + TSP(NN+2-opt)",
            "n_locations": n, "tour": tour, "tour_labels": tour_labels,
            "total_distance": dist, "n_vehicles": 1,
            "summary": f"最短路径: {n}个地点, 总距离={dist}km, 1辆车"
        }
        print(f"     TSP: {dist} total distance, 1 vehicle")
    else:
        # VRP: TSP tour → split by capacity
        _, tour = _tsp_solve(D, n)
        routes, i = [], 0
        while i < len(tour) - 1:
            route, load = [0], 0
            while i < len(tour) - 1 and load + demands[tour[i+1]] <= capacity:
                route.append(tour[i+1]); load += demands[tour[i+1]]; i += 1
            route.append(0); routes.append(route)
            if len(route) == 2: i += 1  # single stop = force advance

        total = 0; details = []
        for ri, r in enumerate(routes):
            rd = sum(D[r[j]][r[j+1]] for j in range(len(r)-1))
            rl = sum(demands[j] for j in r if j != 0)
            total += rd
            details.append({"route": ri+1, "path": [labels[j] for j in r],
                           "distance": round(rd, 1), "load": round(rl, 0)})
            print(f"     Route {ri+1}: {len(r)-2} stops, dist={rd:.1f}, load={rl:.0f}/{capacity:.0f}")

        results[f"sub_{sp['id']}"] = {
            "method": "Floyd-Warshall + TSP + VRP split",
            "n_locations": n, "n_vehicles": len(routes), "routes": details,
            "total_distance": round(total, 1), "vehicle_capacity": capacity,
            "summary": f"VRP: {n}地点->{len(routes)}辆车, 总距离={total:.1f}"
        }
        print(f"     VRP: {len(routes)} vehicles, total distance={total:.1f}")


def _solve_routing_from_coords(sp, numeric, df, results, gs, n, sp_text):
    """坐标数据 → 欧氏距离 → TSP + 2-opt"""
    # Find coordinate columns
    coord_cols = []
    for col in numeric.columns:
        cl = str(col).lower()
        if any(kw in cl for kw in ['x', 'y', '坐标', '经度', '纬度', 'lat', 'lon', 'lng']):
            coord_cols.append(col)
    if len(coord_cols) < 2:
        coord_cols = numeric.columns[:2].tolist()

    coords = numeric[coord_cols].values.astype(float)
    print(f"     Routing: {n} locations from coordinates {coord_cols}")

    # Euclidean distance matrix
    D = np.zeros((n, n))
    for i in range(n):
        for j in range(n):
            D[i, j] = np.sqrt(np.sum((coords[i] - coords[j])**2))

    # Capacity and demands
    sp_text = sp.get("title", "") + sp.get("full_text", "")
    cap_match = re.search(r'(\d+)\s*(kg|千克|公斤|吨|t)', sp_text.lower())
    capacity = float(cap_match.group(1)) if cap_match else float('inf')
    if cap_match and cap_match.group(2) in ('吨', 't'):
        capacity *= 1000

    demand_col = None
    for col in df.columns:
        cl = str(col).lower()
        if any(kw in cl for kw in ['需求', '重量', 'demand', 'weight', 'load']):
            demand_col = col; break
    demands = df[demand_col].values.astype(float).tolist() if demand_col else None

    labels = df.iloc[:, 0].tolist() if not pd.api.types.is_numeric_dtype(df.iloc[:, 0]) else \
             [f"地点{i+1}" for i in range(n)]

    # Solve TSP
    dist, tour = _tsp_solve(D, n)
    tour_labels = [labels[i] for i in tour if i < len(labels)]
    results[f"sub_{sp['id']}"] = {
        "method": "Euclidean TSP (NN+2-opt)",
        "n_locations": n, "tour": tour, "tour_labels": tour_labels,
        "total_distance": dist, "n_vehicles": 1,
        "summary": f"最短路径: {n}个地点, 总距离={dist}, 1辆车"
    }
    print(f"     TSP: {dist} total distance")


def _make_routing_figure(val, key, fig_dir):
    """生成路径优化图：有意义的TSP路线图 — 每问不同"""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from mathmodel.visualization.styles import despine, get_colors

    tour = val.get("tour", [])
    total_dist = val.get("total_distance", 0)
    method = val.get("method", "")
    n_locations = val.get("n_locations", len(tour)-1 if tour else 0)

    if not tour:
        return

    colors = get_colors(5)
    n = len(tour)

    # Create a meaningful route diagram with distance info
    fig, ax = plt.subplots(figsize=(max(8, n*0.5), 4.5))

    # Generate Y positions that show the route flow
    xs = list(range(n))
    # Use sine wave to make route visually interesting
    ys = [np.sin(i * 2 * np.pi / (n-1)) * 1.5 for i in range(n)]

    # Draw the route with direction arrows
    for i in range(n - 1):
        ax.annotate('', xy=(xs[i+1], ys[i+1]), xytext=(xs[i], ys[i]),
                   arrowprops=dict(arrowstyle='->', color=colors[0], lw=2,
                                  connectionstyle='arc3,rad=0.1'))

    # Draw nodes
    for i, node in enumerate(tour):
        is_start = (i == 0)
        ax.scatter(xs[i], ys[i], s=200 if is_start else 120,
                  c=colors[2] if is_start else 'white',
                  edgecolors=colors[0], linewidth=2 if is_start else 1.5,
                  zorder=5)
        label = f'{node+1}'
        if is_start: label += ' (Start)'
        ax.annotate(label, (xs[i], ys[i] + 0.35), fontsize=8, ha='center',
                   fontweight='bold' if is_start else 'normal')

    # Distance info
    q_id = key.replace('sub_', 'Q')
    ax.set_title(f"{q_id}: Optimal Route — {total_dist} km, {n_locations} locations, {method[:30]}",
                fontsize=11, fontweight='bold', loc='left')
    ax.set_xlabel('Visit Sequence'); ax.set_yticks([])
    ax.set_xlim(-0.5, n - 0.5)
    despine(ax); ax.grid(alpha=0.2, axis='x', linestyle=":")

    fig.tight_layout()
    path = Path(fig_dir) / f"{key}_route.png"
    fig.savefig(str(path), dpi=300, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    print(f"  [{key}] Route figure: {total_dist}km")


def _find_df(data_files, ptype):
    for k, v in data_files.items():
        if k.endswith("_norm"): continue
        if v.select_dtypes(include=np.number).shape[1] >= 2: return v
    return list(data_files.values())[0] if data_files else None


def _serialize(obj):
    if isinstance(obj, dict): return {str(k): _serialize(v) for k, v in obj.items()}
    if isinstance(obj, list): return [_serialize(v) for v in obj]
    if isinstance(obj, (np.integer,)): return int(obj)
    if isinstance(obj, (np.floating,)): return float(obj)
    if isinstance(obj, np.ndarray): return obj.tolist()
    return obj


if __name__ == "__main__":
    main()
